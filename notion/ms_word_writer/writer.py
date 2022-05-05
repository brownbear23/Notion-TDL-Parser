from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from notion.ms_word_writer.util import time_util
from notion.constants import DocConstants
from docx.oxml.ns import qn

def tdl_writer(template_docx_dir, doc_name, board_list):
    document = Document(template_docx_dir)

    _add_paragraph(document, "     ↳  " + time_util.get_current_day_weekday(), 15, True)

    current_tdl_field = ""
    for board in board_list:
        if not current_tdl_field == board.tdl_field:
            _add_paragraph(document, "", 22, True)
            _add_paragraph(document, board.tdl_field, 22, True)
            current_tdl_field = board.tdl_field
        _add_task_bullet_list(document, board.title, board.due_date, 20, board.urgency)
        _add_bullet_list(document, board.root_bulleted_block, 20, 2)

    document.save(doc_name + '.docx')




def _add_paragraph(document, content, size, bold):
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.line_spacing = DocConstants.line_space
    paragraph.paragraph_format.space_before = DocConstants.paragraph_space_before
    paragraph.paragraph_format.space_after = DocConstants.paragraph_space_after

    run = paragraph.add_run()
    run.text = content
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = DocConstants.font_name


def _add_task_bullet_list(document, title, due_date, font_size, urgent_lv):
    paragraph = document.add_paragraph(style='Body Text')

    paragraph.paragraph_format.left_indent = Pt(DocConstants.left_indent_num)

    paragraph.paragraph_format.line_spacing = DocConstants.line_space
    paragraph.paragraph_format.space_before = DocConstants.paragraph_space_before
    paragraph.paragraph_format.space_after = DocConstants.paragraph_space_after

    highlight_color = None
    if urgent_lv == 1:
        highlight_color = WD_COLOR_INDEX.RED
    elif urgent_lv == 2:
        highlight_color = WD_COLOR_INDEX.YELLOW
    elif urgent_lv == 3:
        highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    run = paragraph.add_run()
    run.font.size = Pt(font_size - 5)
    run.font.highlight_color = highlight_color

    run.text = "➤   "


    run = paragraph.add_run()
    if "ASAP" in title:
        run.font.bold = True
    else:
        run.font.bold = False

    run.font.size = Pt(font_size)
    run.font.name = DocConstants.font_name
    run.font.highlight_color = highlight_color
    run.text = title
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Dotum')

    if due_date is not None:
        run = paragraph.add_run()
        run.bold = True
        run.italic = True
        run.font.size = Pt(font_size)
        run.font.name = DocConstants.font_name
        run.font.highlight_color = highlight_color
        date_time_obj = time_util.get_date_time(due_date)
        date_time_str = date_time_obj.strftime("%b %d (%a) %H:%M")

        if "00:00" in date_time_str:
            date_time_str = date_time_str.replace("00:00", "")

        run.text = " - " + date_time_str





def _add_bullet_list(document, bulleted_block, font_size, indent_lv):
    if not bulleted_block.root:
        paragraph = document.add_paragraph()

        paragraph.paragraph_format.line_spacing = DocConstants.line_space
        paragraph.paragraph_format.space_before = DocConstants.paragraph_space_before
        paragraph.paragraph_format.space_after = DocConstants.paragraph_space_after

        indent_pt = indent_lv * (DocConstants.left_indent_num + 3)
        paragraph.paragraph_format.left_indent = Pt(indent_pt)

        run = paragraph.add_run()
        run.font.size = Pt(font_size)
        run.font.name = DocConstants.font_name
        run.text = "•   "

        run = paragraph.add_run()
        run.font.size = Pt(font_size)
        run.font.name = DocConstants.font_name

        run.text = bulleted_block.text

        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Dotum')

        run.font.bold = bulleted_block.bold
        run.font.italic = bulleted_block.italic
        run.font.strike = bulleted_block.strikethrough
        run.font.underline = bulleted_block.underline

        indent_lv += 1

    if bulleted_block.has_children:
        for child_bulleted_block in bulleted_block.children:
            _add_bullet_list(document, child_bulleted_block, font_size, indent_lv)








if __name__ == "__main__":
    pass
